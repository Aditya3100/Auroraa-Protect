# app/services/watermark/document/doc_extractor.py

import fitz
from docx import Document


# -----------------------------
# Unicode channel
# -----------------------------

Z0 = "\u200B"   # 0
Z1 = "\u200C"   # 1


# -----------------------------
# Parameters
# -----------------------------

REPEAT = 6
MAGIC = "WM1|"

MAX_PAYLOAD = 4096


# -----------------------------
# Helpers
# -----------------------------

def _majority(bits):

    out = []

    for i in range(0, len(bits), REPEAT):

        block = bits[i:i + REPEAT]

        if len(block) < REPEAT:
            break

        out.append(1 if sum(block) > REPEAT / 2 else 0)

    return out


def _bits_to_bytes(bits):

    buf = bytearray()

    for i in range(0, len(bits), 8):

        chunk = bits[i:i + 8]

        if len(chunk) < 8:
            break

        val = 0

        for b in chunk:
            val = (val << 1) | b

        buf.append(val)

    return bytes(buf)


def _parse_frame(text: str):

    start = text.find(MAGIC)

    if start < 0:
        return None

    rest = text[start + len(MAGIC):]

    sep = rest.find("|")

    if sep < 0:
        return None

    length = rest[:sep]

    if not length.isdigit():
        return None

    length = int(length)

    if length <= 0 or length > MAX_PAYLOAD:
        return None

    payload = rest[sep + 1: sep + 1 + length]

    if len(payload) != length:
        return None

    return payload


# -----------------------------
# Unicode channel
# -----------------------------

def _extract_unicode(text):

    bits = []

    for ch in text:

        if ch == Z0:
            bits.append(0)

        elif ch == Z1:
            bits.append(1)

    return bits


# -----------------------------
# Layout channel (PDF)
# -----------------------------

def _extract_pdf_layout(page):

    bits = []

    blocks = page.get_text("dict")["blocks"]

    for b in blocks:

        if b["type"] != 0:
            continue

        for l in b["lines"]:

            y = l["bbox"][1]

            # Quantize baseline
            base = round(y)

            bits.append(1 if y > base else 0)

    return bits


# -----------------------------
# Layout channel (DOCX)
# -----------------------------

def _extract_docx_layout(doc):

    bits = []

    for p in doc.paragraphs:

        sp = p.paragraph_format.line_spacing

        if not sp:
            continue

        bits.append(1 if sp > 1.05 else 0)

    return bits


# -----------------------------
# Merge channels
# -----------------------------

def _merge(a, b):

    n = min(len(a), len(b))

    out = []

    for i in range(n):

        out.append(1 if a[i] + b[i] >= 1 else 0)

    return out


# -----------------------------
# Decode
# -----------------------------

def _decode(bits):

    clean = _majority(bits)

    data = _bits_to_bytes(clean)

    try:
        txt = data.decode("utf-8", errors="ignore")
    except Exception:
        return None

    return _parse_frame(txt)


# -----------------------------
# PDF
# -----------------------------

def extract_pdf_bits(path, max_bits=4096):

    doc = fitz.open(path)

    unicode_bits = []
    layout_bits = []

    for page in doc:

        text = page.get_text("text")

        unicode_bits += _extract_unicode(text)

        layout_bits += _extract_pdf_layout(page)

    merged = _merge(unicode_bits, layout_bits)

    return _decode(merged)


# -----------------------------
# DOCX
# -----------------------------

def extract_docx_bits(path, max_bits=4096):

    doc = Document(path)

    unicode_bits = []
    layout_bits = _extract_docx_layout(doc)

    for p in doc.paragraphs:

        unicode_bits += _extract_unicode(p.text)

    merged = _merge(unicode_bits, layout_bits)

    return _decode(merged)
